import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip('\n'))
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Handle Tables
        if "|" in line:
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator row (e.g., |---|)
            if re.match(r'^[\s|:-]+$', stripped):
                continue
            
            # Split and clean cells. Ensure it's not a False positive.
            parts = line.split("|")
            # Usually tables start and end with |, so the first and last parts are empty.
            if parts[0].strip() == "": parts = parts[1:]
            if parts and parts[-1].strip() == "": parts = parts[:-1]
            
            cells = [c.strip() for c in parts]
            if cells:
                table_data.append(cells)
            continue
        else:
            if in_table:
                # Render the table
                if table_data:
                    columns = max(len(row) for row in table_data)
                    table = doc.add_table(rows=0, cols=columns)
                    table.style = 'Table Grid'
                    for row_cells in table_data:
                        row = table.add_row()
                        for i, cell_text in enumerate(row_cells):
                            if i < columns:
                                row.cells[i].text = cell_text
                in_table = False
                table_data = []

        # Handle Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        
        # Handle List Items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            doc.add_paragraph(content, style='List Number')
        
        # Handle Body Text
        elif stripped:
            # Process simple bold/italic markers by stripping them (simple version)
            text = stripped.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            doc.add_paragraph(text)
        else:
            # Empty line -> new paragraph for spacing
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

if __name__ == "__main__":
    md_file = "TECHNICAL_DOCUMENTATION.md"
    docx_file = "PCIC_RO10_Technical_Documentation_v1.4.1.docx"
    markdown_to_docx(md_file, docx_file)
